"""施工管理工具集：进度初始化/甘特图/文件填写/到期提醒

功能：
  - construction_init_schedule  根据项目类型自动生成施工进度计划
  - construction_list_schedule  查看施工进度表
  - construction_update_milestone 更新节点进度
  - construction_gantt          生成甘特图（HTML）
  - construction_fill_form      填写施工文件模板（docx）
  - construction_overview       施工全景视图
"""

import json
import logging
import os
import re
import shutil
import tempfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# ── 路径 ─────────────────────────────────────────────
_SCHEDULES_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "construction" / "schedules"
_TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "construction" / "templates"

# ── SpaceType 映射 ──────────────────────────────────
SPACE_TYPE_LABEL = {
    "residential": "私宅",
    "restaurant": "餐饮空间",
    "hotel": "酒店/民宿",
    "exhibition": "展厅",
    "retail": "门店",
}

# ── 工具定义 ────────────────────────────────────────

TOOL_DEFS = [
    {
        "name": "construction_init_schedule",
        "description": "为项目创建施工进度计划表。根据项目类型自动生成标准施工节点（拆改→水电→防水→瓦木→油漆→安装→验收）和计划日期。设计师说'创建施工计划''生成进度表''排工期'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
                "start_date": {"type": "string", "description": "开工日期，格式 YYYY-MM-DD，默认今天"},
                "space_type": {
                    "type": "string",
                    "enum": ["residential", "restaurant", "hotel", "exhibition", "retail"],
                    "description": "项目类型（可选，默认从客户数据推断）",
                },
            },
            "required": ["client_name"],
        },
    },
    {
        "name": "construction_list_schedule",
        "description": "查看项目的施工进度表。显示所有节点的计划日期、实际日期、状态和执行人。设计师说'看看XX的进度''施工到哪了''进度表'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
            },
            "required": ["client_name"],
        },
    },
    {
        "name": "construction_update_milestone",
        "description": "更新施工节点的进度状态（未开始→进行中→已完成）。设计师说'XX完工了''XX开始做了''XX通过了'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
                "milestone_name": {"type": "string", "description": "施工节点名称，如'水电改造''防水工程'"},
                "status": {
                    "type": "string",
                    "enum": ["进行中", "已完成"],
                    "description": "新状态",
                },
            },
            "required": ["client_name", "milestone_name", "status"],
        },
    },
    {
        "name": "construction_gantt",
        "description": "生成项目的甘特图（HTML格式），直观展示施工进度时间线。设计师说'看甘特图''时间线''可视化进度'时使用。返回HTML文件路径。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
            },
            "required": ["client_name"],
        },
    },
    {
        "name": "construction_fill_form",
        "description": "填写施工文件模板。从 data/construction/templates/ 目录读取 docx 模板，用项目数据替换 {{变量}} 占位符，生成填好的文档。设计师说'填写施工日志''生成验收单'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
                "template_name": {"type": "string", "description": "模板文件名（含扩展名），如'施工日志.docx'"},
            },
            "required": ["client_name", "template_name"],
        },
    },
    {
        "name": "construction_overview",
        "description": "查看项目的施工全景视图：当前进度百分比、各节点状态、即将到期任务、已延期任务。设计师说'看看XX的施工全景''整体施工情况'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
            },
            "required": ["client_name"],
        },
    },
    {
        "name": "construction_dashboard",
        "description": "多工地总览仪表盘。一屏查看所有工地的进度状态、延期情况、今日待办。设计师说'看看所有工地''全局施工情况''哪些工地延期了'时使用。无需参数。",
        "parameters": {
            "type": "object",
            "properties": {},
            "required": [],
        },
    },
    {
        "name": "construction_photo_link",
        "description": "将最近一张未关联的施工照片关联到项目的指定节点。设计师说'把照片关联到XX的XX''这张照片是XX的XX'时使用。",
        "parameters": {
            "type": "object",
            "properties": {
                "client_name": {"type": "string", "description": "客户姓名/项目名称"},
                "milestone_name": {"type": "string", "description": "施工节点名称，如防水工程"},
            },
            "required": ["client_name", "milestone_name"],
        },
    },
]


# ════════════════════════════════════════════════════════════
# 数据存储
# ════════════════════════════════════════════════════════════

_CONSTRUCTION_DB: dict[str, dict] = {}  # client_name → schedule


def _db_path() -> Path:
    return Path(__file__).resolve().parent.parent.parent / "data" / "construction" / "schedules_db.json"


def _load_db():
    global _CONSTRUCTION_DB
    p = _db_path()
    if p.exists():
        try:
            _CONSTRUCTION_DB = json.loads(p.read_text("utf-8"))
        except (json.JSONDecodeError, OSError):
            _CONSTRUCTION_DB = {}


def _save_db():
    p = _db_path()
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(_CONSTRUCTION_DB, ensure_ascii=False, indent=2), encoding="utf-8")


_load_db()


def _get_client_data(client_name: str) -> dict:
    """从本地客户数据库查询客户信息"""
    try:
        from max_system.tools.clientmgr_tools import _get_clients_db
        db = _get_clients_db()
        for c in db.values():
            if c.get("name", "").lower() == client_name.lower():
                return c
    except Exception:
        pass
    return {}


def _ensure_schedule(client_name: str) -> dict | None:
    """确保项目已有施工计划"""
    for key in _CONSTRUCTION_DB:
        if key.lower() == client_name.lower():
            return _CONSTRUCTION_DB[key]
    return None


# ════════════════════════════════════════════════════════════
# 1. 初始化施工计划
# ════════════════════════════════════════════════════════════

def _load_schedule_template(space_type: str) -> dict | None:
    path = _SCHEDULES_DIR / f"{space_type}.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text("utf-8"))
    except (json.JSONDecodeError, OSError):
        return None


def _calc_dates(template: dict, start_date: datetime) -> list[dict]:
    """根据模板里程碑和依赖关系计算每个节点的计划日期"""
    milestones = []
    for i, m in enumerate(template["milestones"]):
        if not m["deps"]:
            # 无依赖 → 从开工日期开始
            start = start_date
        else:
            # 取所有依赖节点中完成最晚的那个
            latest_end = start_date
            for dep_idx in m["deps"]:
                if dep_idx < len(milestones):
                    dep = milestones[dep_idx]
                    dep_end = dep["plan_date"] + timedelta(days=dep["days"])
                    if dep_end > latest_end:
                        latest_end = dep_end
            start = latest_end

        plan_date = start
        end_date = start + timedelta(days=m["days"] - 1)

        milestones.append({
            "index": i,
            "phase": m["phase"],
            "name": m["name"],
            "days": m["days"],
            "status": "未开始",
            "plan_date": plan_date.strftime("%Y-%m-%d"),
            "end_date": end_date.strftime("%Y-%m-%d"),
            "actual_date": "",
            "tasks": m.get("tasks", []),
            "workers": m.get("workers", []),
        })
    return milestones


async def construction_init_schedule(args: dict) -> dict:
    client_name = args.get("client_name", "")
    start_str = args.get("start_date", "")
    space_type = args.get("space_type", "")

    if not client_name:
        return {"content": [{"type": "text", "text": "请指定客户姓名或项目名称"}]}

    # 从客户数据获取 space_type
    client = _get_client_data(client_name)
    if not space_type:
        from max_system.integrations.ark_bridge import infer_space_type
        space_type = infer_space_type(client)
    if space_type not in SPACE_TYPE_LABEL:
        space_type = "residential"

    # 加载模板
    template = _load_schedule_template(space_type)
    if not template:
        return {"content": [{"type": "text", "text": f"未找到 {space_type} 类型的施工模板"}]}

    # 计算日期
    start_date = datetime.now()
    if start_str:
        try:
            start_date = datetime.strptime(start_str, "%Y-%m-%d")
        except ValueError:
            pass

    milestones = _calc_dates(template, start_date)

    schedule = {
        "client_name": client_name,
        "space_type": space_type,
        "label": template["label"],
        "total_days": template["totalDays"],
        "start_date": start_date.strftime("%Y-%m-%d"),
        "created_at": datetime.now().isoformat(),
        "milestones": milestones,
        "updated_at": datetime.now().isoformat(),
    }

    # 保存
    _CONSTRUCTION_DB[client_name] = schedule
    _save_db()

    # 生成任务（同步到飞书）
    task_count = 0
    try:
        from max_system.tools.feishu_tools import _get_api_client
        client_api = _get_api_client()
        if client_api:
            for m in milestones:
                task_title = f"[施工] {client_name} - {m['name']}"
                due = m["end_date"] + "T18:00:00"
                try:
                    await client_api.create_task(
                        summary=task_title,
                        due_at=due,
                        description="\n".join(f"- {t}" for t in m["tasks"]),
                    )
                    task_count += 1
                except Exception:
                    pass
    except Exception:
        pass

    lines = [
        f"📋 **{client_name}** 施工计划已创建",
        f"类型：{template['label']} | 工期：{template['totalDays']}天 | 开工：{start_date.strftime('%Y-%m-%d')}",
        "",
        "**施工节点：**",
    ]
    phase_map = {1: "第一阶段", 2: "第二阶段", 3: "第三阶段"}
    current_phase = 0
    for m in milestones:
        if m["phase"] != current_phase:
            current_phase = m["phase"]
            lines.append(f"\n**{phase_map.get(current_phase, f'阶段{current_phase}')}**")
        lines.append(f"  {m['name']}  {m['plan_date']}→{m['end_date']} ({m['days']}天)  {m['status']}")

    if task_count > 0:
        lines.append(f"\n已同步 {task_count} 个节点到飞书任务")

    lines.append(f"\n💡 后续可用命令：")
    lines.append("  • `查看XX的进度` — 查看进度表")
    lines.append("  • `XX水电改造完成了` — 更新节点状态")
    lines.append("  • `看XX的甘特图` — 可视化时间线")

    return {"content": [{"type": "text", "text": "\n".join(lines)}]}


# ════════════════════════════════════════════════════════════
# 2. 查看施工计划
# ════════════════════════════════════════════════════════════

async def construction_list_schedule(args: dict) -> dict:
    client_name = args.get("client_name", "")
    if not client_name:
        return {"content": [{"type": "text", "text": "请指定客户姓名"}]}

    schedule = _ensure_schedule(client_name)
    if not schedule:
        return {"content": [{"type": "text", "text": f"未找到 {client_name} 的施工计划。请先使用「创建施工计划」初始化。"}]}

    now = datetime.now().date()
    milestones = schedule["milestones"]

    # 统计数据
    total = len(milestones)
    done = sum(1 for m in milestones if m["status"] == "已完成")
    progress = round(done / total * 100) if total > 0 else 0
    overdue = sum(1 for m in milestones if m["status"] != "已完成" and m["end_date"] < now.strftime("%Y-%m-%d"))

    lines = [
        f"📋 **{client_name}** 施工进度",
        f"类型：{schedule['label']} | 开工：{schedule['start_date']} | 整体进度：{progress}% ({done}/{total})",
        "",
    ]
    if overdue > 0:
        lines.append(f"⚠️  **{overdue} 个节点已延期！**")
        for m in milestones:
            if m["status"] != "已完成" and m["end_date"] < now.strftime("%Y-%m-%d"):
                lines.append(f"  ❗ {m['name']} 计划 {m['end_date']}，已延期 {(now - datetime.strptime(m['end_date'], '%Y-%m-%d').date()).days} 天")
        lines.append("")

    # 表格
    lines.append("| 节点 | 计划 | 实际 | 状态 | 执行人 |")
    lines.append("|------|------|------|------|--------|")
    for m in milestones:
        status_icon = {"未开始": "⬜", "进行中": "🔄", "已完成": "✅"}.get(m["status"], "⬜")
        actual = m["actual_date"] or "-"
        workers = ", ".join(m["workers"]) if m["workers"] else "-"
        lines.append(f"| {status_icon} {m['name']} | {m['plan_date']}→{m['end_date']} | {actual} | {m['status']} | {workers} |")

    return {"content": [{"type": "text", "text": "\n".join(lines)}]}


# ════════════════════════════════════════════════════════════
# 3. 更新节点进度
# ════════════════════════════════════════════════════════════

async def construction_update_milestone(args: dict) -> dict:
    client_name = args.get("client_name", "")
    milestone_name = args.get("milestone_name", "")
    new_status = args.get("status", "")

    if not all([client_name, milestone_name, new_status]):
        return {"content": [{"type": "text", "text": "请提供客户姓名、节点名称和新状态"}]}

    schedule = _ensure_schedule(client_name)
    if not schedule:
        return {"content": [{"type": "text", "text": f"未找到 {client_name} 的施工计划"}]}

    found = False
    for m in schedule["milestones"]:
        if m["name"] == milestone_name:
            old_status = m["status"]
            m["status"] = new_status
            if new_status == "已完成" and not m["actual_date"]:
                m["actual_date"] = datetime.now().strftime("%Y-%m-%d")
            # 级联更新后续节点
            if new_status == "已完成":
                _cascade_update(schedule, m)
            found = True
            break

    if not found:
        return {"content": [{"type": "text", "text": f"未找到节点「{milestone_name}」，可用节点：{', '.join(m['name'] for m in schedule['milestones'])}"}]}

    schedule["updated_at"] = datetime.now().isoformat()
    _save_db()

    # 节点完成 → 回款提醒
    payment_tip = ""
    if new_status == "已完成":
        try:
            total = len(schedule["milestones"])
            idx = next((m["index"] for m in schedule["milestones"] if m["name"] == milestone_name), 0)
            ratio = round((idx + 1) / total * 100)
            payment_tip = (
                f"\n\n💰 **回款提醒**：「{milestone_name}」已完成，"
                f"建议按合同催收该阶段进度款。当前进度 {ratio}%（{idx+1}/{total} 节点）"
            )
        except Exception:
            pass

    text = f"✅ {milestone_name} → {new_status}{payment_tip}"
    return {"content": [{"type": "text", "text": text}]}


def _cascade_update(schedule: dict, completed_milestone: dict):
    """节点完成后，如果有依赖此节点的后续节点且状态是'未开始'，且其所有依赖都已完成，会自动更新为'进行中'"""
    idx = completed_milestone["index"]
    for m in schedule["milestones"]:
        if m["status"] == "未开始" and idx in m.get("deps", []):
            # 检查是否所有依赖都已完成
            all_deps_done = all(
                schedule["milestones"][d]["status"] == "已完成"
                for d in m.get("deps", [])
            )
            if all_deps_done:
                m["status"] = "进行中"


# ════════════════════════════════════════════════════════════
# 4. 甘特图（HTML）
# ════════════════════════════════════════════════════════════

async def construction_gantt(args: dict) -> dict:
    client_name = args.get("client_name", "")
    if not client_name:
        return {"content": [{"type": "text", "text": "请指定客户姓名"}]}

    schedule = _ensure_schedule(client_name)
    if not schedule:
        return {"content": [{"type": "text", "text": f"未找到 {client_name} 的施工计划"}]}

    milestones = schedule["milestones"]
    if not milestones:
        return {"content": [{"type": "text", "text": "施工计划为空"}]}

    # 计算时间范围
    all_dates = []
    for m in milestones:
        all_dates.append(datetime.strptime(m["plan_date"], "%Y-%m-%d"))
        all_dates.append(datetime.strptime(m["end_date"], "%Y-%m-%d"))
    start = min(all_dates)
    end = max(all_dates)
    total_days = (end - start).days + 1

    # 每周一标记
    week_starts = []
    d = start
    while d <= end:
        if d.weekday() == 0:  # Monday
            week_starts.append(d)
        d += timedelta(days=1)

    status_colors = {"未开始": "#94a3b8", "进行中": "#3b82f6", "已完成": "#22c55e"}
    phase_colors = {1: ["#eff6ff", "#bfdbfe"], 2: ["#f0fdf4", "#bbf7d0"], 3: ["#fefce8", "#fde68a"]}

    # 计算每个条的偏移
    bar_info = []
    for m in milestones:
        ps = datetime.strptime(m["plan_date"], "%Y-%m-%d")
        pe = datetime.strptime(m["end_date"], "%Y-%m-%d")
        left = (ps - start).days
        width = max((pe - ps).days + 1, 1)
        bar_info.append({
            "name": m["name"],
            "left": left,
            "width": width,
            "status": m["status"],
            "phase": m["phase"],
            "plan_date": m["plan_date"],
            "end_date": m["end_date"],
            "actual_date": m.get("actual_date", ""),
            "tasks": m.get("tasks", []),
        })

    html = _gantt_html(client_name, schedule, bar_info, start, total_days, week_starts, phase_colors, status_colors)

    output_dir = Path(__file__).resolve().parent.parent.parent / "data" / "construction"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"gantt_{client_name}.html"
    # 清理文件名
    safe_name = re.sub(r'[\\/*?:"<>|]', '_', client_name)
    output_path = output_dir / f"gantt_{safe_name}.html"
    output_path.write_text(html, encoding="utf-8")

    return {"content": [{"type": "text", "text": f"甘特图已生成：{output_path}\n请在浏览器中打开查看。"}]}


def _gantt_html(client_name, schedule, bars, start, total_days, week_starts, phase_colors, status_colors):
    day_w = 28  # 每天像素宽度
    row_h = 36
    header_h = 50
    label_w = 140
    padding = 12

    rows = []
    for i, b in enumerate(bars):
        y = header_h + i * row_h
        phase = b["phase"]
        bg = phase_colors.get(phase, ["#fff", "#fff"])[0]
        bar_color = status_colors.get(b["status"], "#94a3b8")

        extra = ""
        if b["actual_date"]:
            extra = f'<div style="position:absolute;left:{label_w + (b["left"] + b["width"]) * day_w + 8}px;top:{y}px;font-size:11px;color:#64748b;line-height:30px">✔ {b["actual_date"]}</div>'

        rows.append(f"""
        <div style="position:absolute;top:{y}px;left:0;right:0;height:{row_h}px;background:{bg};border-bottom:1px solid #e2e8f0">
          <div style="position:absolute;left:{padding}px;top:0;width:{label_w}px;height:100%;line-height:{row_h}px;font-size:13px;color:#1e293b;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;font-weight:{'600' if b['status']=='已完成' else '400'}">{b['name']}</div>
          <div style="position:absolute;left:{label_w + b['left'] * day_w}px;top:{row_h//2 - 8}px;width:{b['width'] * day_w}px;height:16px;background:{bar_color};border-radius:8px;opacity:{0.9 if b['status']=='已完成' else 0.7};transition:all .3s" title="{b['plan_date']} → {b['end_date']}"></div>
          {extra}
        </div>""")

    # 时间轴头部
    headers = []
    for w in week_starts:
        day_offset = (w - start).days
        x = label_w + day_offset * day_w
        headers.append(f'<div style="position:absolute;left:{x}px;top:0;font-size:11px;color:#64748b;line-height:{header_h}px;width:80px">{w.strftime("%m/%d")}</div>')

    # 竖网格线（每周）
    grid_lines = []
    for w in week_starts:
        day_offset = (w - start).days
        x = label_w + day_offset * day_w
        grid_lines.append(f'<div style="position:absolute;left:{x}px;top:{header_h}px;bottom:0;width:1px;background:#e2e8f0"></div>')

    total_w = label_w + total_days * day_w + 40

    # 统计
    done = sum(1 for b in bars if b["status"] == "已完成")
    progress = round(done / len(bars) * 100) if bars else 0
    now = datetime.now().date()
    overdue = sum(1 for b in bars if b["status"] != "已完成" and b["end_date"] < now.strftime("%Y-%m-%d") if datetime.strptime(b["end_date"], "%Y-%m-%d").date() < now)

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>施工甘特图 - {client_name}</title>
<style>
  *{{margin:0;padding:0;box-sizing:border-box}}
  body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:#f8fafc;padding:20px}}
  .header{{margin-bottom:20px}}
  .header h1{{font-size:20px;color:#0f172a}}
  .header .meta{{font-size:13px;color:#64748b;margin-top:4px}}
  .stats{{display:flex;gap:16px;margin:12px 0}}
  .stat{{padding:8px 16px;border-radius:8px;font-size:13px}}
  .stat.progress{{background:#dbeafe;color:#1d4ed8}}
  .stat.done{{background:#dcfce7;color:#16a34a}}
  .stat.overdue{{background:#fee2e2;color:#dc2626}}
  .gantt-wrap{{overflow-x:auto;background:#fff;border-radius:12px;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
  .gantt{{position:relative;min-height:{(len(bars)+1)*row_h + 20}px}}
  .gantt-canvas{{margin-left:{label_w}px;height:{(len(bars))*row_h}px;position:relative;overflow:hidden}}
</style></head><body>
<div class="header">
  <h1>🏗 {client_name} · 施工甘特图</h1>
  <div class="meta">{schedule['label']} | 开工 {schedule['start_date']} | 共 {len(bars)} 个节点</div>
  <div class="stats">
    <div class="stat progress">📊 进度 {progress}%</div>
    <div class="stat done">✅ 完成 {done}/{len(bars)}</div>
    {f'<div class="stat overdue">⚠️ {overdue} 延期</div>' if overdue else ''}
  </div>
</div>
<div class="gantt-wrap"><div class="gantt" style="width:{total_w}px">
  <div style="position:relative;height:{header_h}px;background:#f1f5f9;border-bottom:2px solid #e2e8f0">
    <div style="position:absolute;left:{padding}px;top:0;width:{label_w}px;height:100%;line-height:{header_h}px;font-size:13px;font-weight:600;color:#475569">施工节点</div>
    {''.join(headers)}
  </div>
  {''.join(grid_lines)}
  {''.join(rows)}
</div></div>
<div style="margin-top:12px;font-size:12px;color:#94a3b8">
  <span style="display:inline-block;width:12px;height:12px;background:#22c55e;border-radius:3px;vertical-align:middle;margin-right:4px"></span> 已完成
  <span style="display:inline-block;width:12px;height:12px;background:#3b82f6;border-radius:3px;vertical-align:middle;margin:0 4px 0 12px"></span> 进行中
  <span style="display:inline-block;width:12px;height:12px;background:#94a3b8;border-radius:3px;vertical-align:middle;margin:0 4px 0 12px"></span> 未开始
  <span style="margin-left:12px">生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}</span>
</div>
</body></html>"""


# ════════════════════════════════════════════════════════════
# 5. 填写施工文件模板
# ════════════════════════════════════════════════════════════

async def construction_fill_form(args: dict) -> dict:
    client_name = args.get("client_name", "")
    template_name = args.get("template_name", "")

    if not client_name or not template_name:
        return {"content": [{"type": "text", "text": "请提供客户姓名和模板文件名"}]}

    template_path = _TEMPLATES_DIR / template_name
    if not template_path.exists():
        available = [f.name for f in _TEMPLATES_DIR.iterdir()] if _TEMPLATES_DIR.exists() else []
        msg = f"模板文件不存在: {template_name}"
        if available:
            msg += f"\n可用模板: {', '.join(available)}"
        else:
            msg += "\n请先将 docx 模板放到 data/construction/templates/ 目录"
        return {"content": [{"type": "text", "text": msg}]}

    # 收集项目数据
    variables = _collect_project_vars(client_name)

    # 检查模板中使用了哪些变量
    template_content = template_path.read_bytes()
    used_vars = re.findall(r'\{\{(\w+)\}\}', template_content.decode("utf-8", errors="ignore"))

    # 填充模板
    output_dir = _TEMPLATES_DIR / "_output"
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r'[\\/*?:"<>|]', '_', client_name)
    output_path = output_dir / f"{safe_name}_{template_name}"

    try:
        from docx import Document
        doc = Document(template_path)

        # 替换段落中的变量
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = _replace_vars(run.text, variables)

        # 替换表格中的变量
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = _replace_vars(run.text, variables)

        doc.save(str(output_path))

        # 列出已填写的变量
        filled = [v for v in used_vars if v in variables]
        missing = [v for v in used_vars if v not in variables]

        lines = [
            f"✅ 已生成: {output_path}\n",
            f"客户: {client_name} | 模板: {template_name}",
            f"已填写 {len(filled)} 个变量，{len(missing)} 个未匹配" if missing else f"已填写全部 {len(filled)} 个变量",
        ]
        if missing:
            lines.append(f"\n未匹配变量: {', '.join(missing)}")
            lines.append("请补充这些值到客户资料，或修改模板中的变量名。")

        return {"content": [{"type": "text", "text": "\n".join(lines)}]}

    except ImportError:
        return {"content": [{"type": "text", "text": "python-docx 未安装，无法处理 .docx 文件"}]}
    except Exception as e:
        return {"content": [{"type": "text", "text": f"文件生成失败: {e}"}]}


def _collect_project_vars(client_name: str) -> dict:
    """收集项目相关数据作为模板变量"""
    vars_dict = {
        "日期": datetime.now().strftime("%Y-%m-%d"),
        "年份": str(datetime.now().year),
        "月份": str(datetime.now().month),
    }

    # 客户数据
    client = _get_client_data(client_name)
    if client:
        vars_dict["客户姓名"] = client.get("name", client_name)
        vars_dict["客户电话"] = client.get("phone", client.get("联系方式", ""))
        vars_dict["项目地址"] = client.get("city", "")
        vars_dict["户型"] = client.get("unit_type", "")
        vars_dict["面积"] = str(client.get("area", ""))
        vars_dict["项目类型"] = SPACE_TYPE_LABEL.get(client.get("type", ""), client.get("type", ""))
        vars_dict["预算金额"] = str(client.get("budget", ""))
        vars_dict["备注"] = client.get("remark", "")

    # Profile
    try:
        from max_system.tools.profile_tools import _get_active_profile_mgr
        mgr = _get_active_profile_mgr()
        if mgr:
            pf = mgr.get_all()
            vars_dict["公司名称"] = pf.get("company_name", "")
            vars_dict["公司电话"] = pf.get("company_phone", "")
            vars_dict["设计师"] = pf.get("designer_name", "")
            vars_dict["公司地址"] = pf.get("company_address", "")
            vars_dict["设计风格"] = pf.get("design_style", "")
    except Exception:
        pass

    # 默认值
    vars_dict.setdefault("客户姓名", client_name)
    vars_dict.setdefault("公司名称", "ARK Design")
    vars_dict.setdefault("设计师", "未设置")

    return vars_dict


def _replace_vars(text: str, variables: dict) -> str:
    """替换文本中的 {{变量}}"""
    def _replacer(m):
        key = m.group(1)
        return str(variables.get(key, m.group(0)))
    return re.sub(r'\{\{(\w+)\}\}', _replacer, text)


# ════════════════════════════════════════════════════════════
# 6. 施工全景视图
# ════════════════════════════════════════════════════════════

async def construction_overview(args: dict) -> dict:
    client_name = args.get("client_name", "")
    if not client_name:
        return {"content": [{"type": "text", "text": "请指定客户姓名"}]}

    schedule = _ensure_schedule(client_name)
    if not schedule:
        return {"content": [{"type": "text", "text": f"未找到 {client_name} 的施工计划。"}]}

    milestones = schedule["milestones"]
    now = datetime.now().date()

    total = len(milestones)
    done = sum(1 for m in milestones if m["status"] == "已完成")
    in_progress = sum(1 for m in milestones if m["status"] == "进行中")
    pending = sum(1 for m in milestones if m["status"] == "未开始")
    progress = round(done / total * 100) if total > 0 else 0

    overdue = []
    coming_soon = []
    for m in milestones:
        end = datetime.strptime(m["end_date"], "%Y-%m-%d").date()
        if m["status"] != "已完成":
            if end < now:
                overdue.append(m)
            elif 0 <= (end - now).days <= 3:
                coming_soon.append(m)

    # 进度条
    bar_len = 20
    filled = round(progress / 100 * bar_len)
    bar = "█" * filled + "░" * (bar_len - filled)

    lines = [
        f"🏗 **{client_name}** · 施工全景",
        f"类型：{schedule['label']} | 开工：{schedule['start_date']}",
        f"进度：{bar} {progress}%",
        f"✅ 已完成 {done} | 🔄 进行中 {in_progress} | ⬜ 待开始 {pending}",
        "",
    ]

    if overdue:
        lines.append(f"🚨 **延期提醒 ({len(overdue)}个）：**")
        for m in overdue:
            days = (now - datetime.strptime(m["end_date"], "%Y-%m-%d").date()).days
            lines.append(f"  ❗ {m['name']} — 已延期 {days} 天（原计划 {m['end_date']}）")

    if coming_soon:
        lines.append(f"\n⏰ **即将到期（{len(coming_soon)}个）：**")
        for m in coming_soon:
            days = (datetime.strptime(m["end_date"], "%Y-%m-%d").date() - now).days
            lines.append(f"  🔔 {m['name']} — {days} 天后到期（{m['end_date']}）")

    lines.append(f"\n💡 可用命令：")
    lines.append("  • `更新XX节点为已完成` — 更新进度")
    lines.append("  • `看XX的甘特图` — 查看时间线")
    lines.append("  • `填写XX的施工日志` — 生成施工文件")

    return {"content": [{"type": "text", "text": "\n".join(lines)}]}


# ════════════════════════════════════════════════════════════
# 6. 多工地总览仪表盘
# ════════════════════════════════════════════════════════════

async def construction_dashboard(args: dict) -> dict:
    """查看所有工地的施工全景仪表盘"""
    _load_db()
    now = datetime.now().date()
    all_schedules = list(_CONSTRUCTION_DB.values())

    if not all_schedules:
        return {"content": [{"type": "text", "text": "当前没有在建项目"}]}

    total = len(all_schedules)
    lines = [
        f"🏗 **多工地总览 · {total} 个在建项目**",
        f"更新：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
    ]

    overdue_global = []
    due_today_global = []

    for s in all_schedules:
        client = s["client_name"]
        milestones = s["milestones"]
        if not milestones:
            continue

        done = sum(1 for m in milestones if m["status"] == "已完成")
        ip = sum(1 for m in milestones if m["status"] == "进行中")
        total_m = len(milestones)
        pct = round(done / total_m * 100)

        # 延期检测
        has_od = False
        for m in milestones:
            if m["status"] == "已完成":
                continue
            try:
                ed = datetime.strptime(m["end_date"], "%Y-%m-%d").date()
                if ed < now:
                    has_od = True
                    overdue_global.append(f"{client}·{m['name']}延{(now-ed).days}天")
                elif ed == now:
                    due_today_global.append(f"{client}·{m['name']}")
            except ValueError:
                pass

        od_mark = " 🚨" if has_od else ""
        ip_label = f" 🔄{ip}项" if ip else ""
        lines.append(
            f"  **{client}**{od_mark}  |  {s['label']}  |  "
            f"█" * (pct // 10) + "░" * (10 - pct // 10) + f" {pct}%"
            f"{ip_label}"
        )

    if due_today_global:
        lines.append(f"\n⚡ **今日到期**")
        for d in due_today_global:
            lines.append(f"  ‣ {d}")

    if overdue_global:
        lines.append(f"\n🚨 **已延期**")
        for d in overdue_global[:10]:
            lines.append(f"  ❗ {d}")
        if len(overdue_global) > 10:
            lines.append(f"  ...还有 {len(overdue_global) - 10} 个")

    lines.append(f"\n💡 查看详情：`看XX的施工全景` 或 `看XX的甘特图`")

    return {"content": [{"type": "text", "text": "\n".join(lines)}]}


# ════════════════════════════════════════════════════════════
# 7. 照片关联施工节点
# ════════════════════════════════════════════════════════════

async def construction_photo_link(args: dict) -> dict:
    """将最近一张未关联的施工照片链接到项目的指定节点"""
    client_name = args.get("client_name", "")
    milestone_name = args.get("milestone_name", "")
    chat_id = args.get("chat_id", "")

    if not client_name or not milestone_name:
        return {"content": [{"type": "text", "text": "请提供项目名称和节点名称"}]}

    # 确保项目存在
    schedule = _ensure_schedule(client_name)
    if not schedule:
        return {"content": [{"type": "text", "text": f"未找到 {client_name} 的施工计划"}]}

    # 验证节点
    milestone_names = [m["name"] for m in schedule["milestones"]]
    if milestone_name not in milestone_names:
        return {"content": [{"type": "text", "text": f"未找到节点「{milestone_name}」，可用节点：{', '.join(milestone_names)}"}]}

    # 读取照片索引
    from pathlib import Path as _Path
    index_path = _Path(__file__).resolve().parent.parent.parent / "data" / "construction" / "photos" / "index.json"
    if not index_path.exists():
        return {"content": [{"type": "text", "text": "没有找到未关联的照片"}]}

    try:
        index_data = json.loads(index_path.read_text("utf-8"))
    except (json.JSONDecodeError, OSError):
        return {"content": [{"type": "text", "text": "读取照片索引失败"}]}

    # 找最近一张未关联且匹配 chat_id 的照片
    photos = index_data.get("photos", [])
    target = None
    for p in reversed(photos):
        if not p.get("project") and not p.get("milestone"):
            if chat_id and p.get("chat_id") == chat_id:
                target = p
                break
            elif not chat_id:
                target = p
                break

    if not target:
        return {"content": [{"type": "text", "text": "没有找到可关联的照片（可能已全部关联）"}]}

    # 更新索引
    target["project"] = client_name
    target["milestone"] = milestone_name
    index_path.write_text(json.dumps(index_data, ensure_ascii=False, indent=2), encoding="utf-8")

    # 物理移动照片到项目文件夹
    photo_path = _Path(__file__).resolve().parent.parent.parent / "data" / "construction" / target["path"]
    if photo_path.exists():
        milestone_dir = _Path(__file__).resolve().parent.parent.parent / "data" / "construction" / "photos" / client_name / milestone_name
        milestone_dir.mkdir(parents=True, exist_ok=True)
        dest = milestone_dir / photo_path.name
        import shutil
        shutil.copy2(str(photo_path), str(dest))

    return {
        "content": [{"type": "text", "text": f"✅ 照片已关联到 {client_name} · {milestone_name}"}]
    }


# ════════════════════════════════════════════════════════════
# 注册
# ════════════════════════════════════════════════════════════

def register_tools(settings) -> list[tuple[str, callable, dict]]:
    handlers = {
        "construction_init_schedule": construction_init_schedule,
        "construction_list_schedule": construction_list_schedule,
        "construction_update_milestone": construction_update_milestone,
        "construction_gantt": construction_gantt,
        "construction_fill_form": construction_fill_form,
        "construction_overview": construction_overview,
        "construction_dashboard": construction_dashboard,
        "construction_photo_link": construction_photo_link,
    }
    return [(d["name"], handlers[d["name"]], d) for d in TOOL_DEFS]
